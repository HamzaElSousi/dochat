import io
import json
import base64
import struct
import zipfile
import pytest
from docx import Document as DocxDocument

VALID_AUTH = {'Authorization': 'Basic ' + base64.b64encode(b'admin:test-password').decode()}

# --- Endpoint integration tests ---

def test_upload_no_auth(client):
    """No Authorization header -> 401."""
    data = {'file': (io.BytesIO(b'hello'), 'test.txt')}
    response = client.post('/admin/ingest/upload', data=data, content_type='multipart/form-data')
    assert response.status_code == 401

def test_upload_file_too_large(client, mocker):
    """File > 10 MB -> 413 before any processing."""
    mocker.patch('app.ingest.embedder.requests.post')  # should never be called
    big_bytes = b'x' * (10 * 1024 * 1024 + 1)
    data = {'file': (io.BytesIO(big_bytes), 'big.txt')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code == 413
    body = json.loads(response.data)
    assert 'error' in body

def test_upload_unsupported_type(client, mocker):
    """Unsupported extension -> 422 with error field."""
    mocker.patch('app.ingest.embedder.requests.post')
    data = {'file': (io.BytesIO(b'PK\x03\x04'), 'archive.zip')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code in (415, 422)
    body = json.loads(response.data)
    assert 'error' in body

def _mock_embed(mocker, n_chunks):
    mock_post = mocker.patch('app.ingest.embedder.requests.post')
    mock_post.return_value.json.return_value = {
        'data': [{'embedding': [0.1] * 1536, 'index': i} for i in range(n_chunks)]
    }
    mock_post.return_value.raise_for_status = lambda: None
    return mock_post

def test_upload_txt_success(client, mocker):
    """Valid TXT upload -> 200, doc_id present, chunk_count > 0, status=ready."""
    text = ('This is a test document. ' * 50).encode('utf-8')
    _mock_embed(mocker, 1)
    data = {'file': (io.BytesIO(text), 'readme.txt')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code == 200
    body = json.loads(response.data)
    assert 'doc_id' in body
    assert body['filename'] == 'readme.txt'
    assert body['chunk_count'] > 0
    assert body['status'] == 'ready'

def test_upload_pdf_success(client, mocker, tmp_path):
    """Valid PDF upload -> 200, doc_id present, chunk_count > 0."""
    mocker.patch('app.ingest.parser.parse_pdf', return_value='Sample PDF text. ' * 30)
    _mock_embed(mocker, 1)
    # Use any bytes -- parse_pdf is mocked
    data = {'file': (io.BytesIO(b'%PDF-1.4 mocked'), 'report.pdf')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code == 200
    body = json.loads(response.data)
    assert body['chunk_count'] > 0
    assert body['status'] == 'ready'

def test_upload_docx_success(client, mocker):
    """Valid DOCX upload -> 200, chunk_count > 0."""
    mocker.patch('app.ingest.parser.parse_docx', return_value='Sample DOCX text. ' * 30)
    _mock_embed(mocker, 1)
    data = {'file': (io.BytesIO(b'PK mocked docx'), 'report.docx')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code == 200
    body = json.loads(response.data)
    assert body['chunk_count'] > 0

def test_upload_corrupt_pdf_returns_422(client, mocker):
    """Corrupt PDF (parse raises ValueError) -> 422 with error, no DB rows."""
    mocker.patch('app.ingest.parser.parse_pdf',
                 side_effect=ValueError('PDF appears to be corrupt'))
    data = {'file': (io.BytesIO(b'not a pdf'), 'bad.pdf')}
    response = client.post('/admin/ingest/upload', data=data,
                           content_type='multipart/form-data', headers=VALID_AUTH)
    assert response.status_code == 422
    body = json.loads(response.data)
    assert 'error' in body
    assert 'Traceback' not in response.data.decode()

def test_upload_duplicate_replaces(client, app, mocker):
    """Same filename uploaded twice -> second replaces first; only one doc row remains."""
    txt = b'Hello world document content for testing duplicates. ' * 20

    for _ in range(2):
        mock_post = mocker.patch('app.ingest.embedder.requests.post')
        mock_post.return_value.json.return_value = {
            'data': [{'embedding': [0.1] * 1536, 'index': 0}]
        }
        mock_post.return_value.raise_for_status = lambda: None
        data = {'file': (io.BytesIO(txt), 'dup.txt')}
        client.post('/admin/ingest/upload', data=data,
                    content_type='multipart/form-data', headers=VALID_AUTH)

    conn = app.config['DB_CONN']
    rows = conn.execute("SELECT COUNT(*) FROM documents WHERE filename='dup.txt'").fetchone()
    assert rows[0] == 1, f"Expected 1 document, got {rows[0]}"

# --- Unit tests for utilities ---

def test_parse_pdf_scanned_empty():
    """parse_pdf with image-only bytes raises ValueError with 'no extractable text'."""
    from app.ingest.parser import parse_pdf
    import unittest.mock as mock
    mock_page = mock.MagicMock()
    mock_page.extract_text.return_value = None
    mock_pdf = mock.MagicMock()
    mock_pdf.__enter__ = mock.MagicMock(return_value=mock_pdf)
    mock_pdf.__exit__ = mock.MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]
    with mock.patch('pdfplumber.open', return_value=mock_pdf):
        with pytest.raises(ValueError, match='no extractable text'):
            parse_pdf(b'%PDF fake')

def test_parse_docx_tables_captured():
    """parse_docx captures text from tables (not just paragraphs)."""
    from app.ingest.parser import parse_docx
    import io as _io
    from docx import Document
    doc = Document()
    doc.add_paragraph("Paragraph text here.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Col A"
    table.cell(0, 1).text = "Col B"
    table.cell(1, 0).text = "Value 1"
    table.cell(1, 1).text = "Value 2"
    buf = _io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert "Col A" in result
    assert "Value 1" in result

def test_chunk_text_token_count():
    """Each chunk from chunk_text() has <= 512 tokens when encoded with cl100k_base."""
    import tiktoken
    from app.ingest.chunker import chunk_text
    long_text = "The quick brown fox jumps over the lazy dog. " * 500
    chunks = chunk_text(long_text)
    assert len(chunks) > 1, "Long text should produce multiple chunks"
    enc = tiktoken.get_encoding("cl100k_base")
    for chunk in chunks:
        token_count = len(enc.encode(chunk))
        assert token_count <= 512, f"Chunk has {token_count} tokens, expected <= 512"

def test_embed_chunks_single_call(mocker):
    """embed_chunks(['a','b','c']) makes exactly 1 requests.post call."""
    mock_post = mocker.patch('app.ingest.embedder.requests.post')
    mock_post.return_value.json.return_value = {
        'data': [{'embedding': [0.1] * 1536, 'index': i} for i in range(3)]
    }
    mock_post.return_value.raise_for_status = lambda: None
    from app.ingest.embedder import embed_chunks
    result = embed_chunks(['chunk one', 'chunk two', 'chunk three'])
    assert mock_post.call_count == 1
    assert len(result) == 3
    assert len(result[0]) == 1536

def test_embed_chunks_subbatch(mocker):
    """embed_chunks with 101 texts makes exactly 2 requests.post calls (sub-batch at 100)."""
    call_count = [0]
    def fake_post(*args, **kwargs):
        call_count[0] += 1
        n = len(kwargs['json']['input'])
        mock_resp = mocker.MagicMock()
        mock_resp.json.return_value = {
            'data': [{'embedding': [0.1] * 1536, 'index': i} for i in range(n)]
        }
        mock_resp.raise_for_status = lambda: None
        return mock_resp
    mocker.patch('app.ingest.embedder.requests.post', side_effect=fake_post)
    from app.ingest.embedder import embed_chunks
    texts = [f'chunk {i}' for i in range(101)]
    result = embed_chunks(texts)
    assert call_count[0] == 2, f"Expected 2 calls, got {call_count[0]}"
    assert len(result) == 101
