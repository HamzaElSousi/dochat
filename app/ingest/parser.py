import io
import os
import zipfile

import pdfplumber
import trafilatura
from pdfminer.pdfdocument import PDFPasswordIncorrect
from pdfminer.pdfparser import PDFSyntaxError
from docx import Document


ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}


def detect_file_type(filename: str) -> str:
    """Return filetype string ('pdf','docx','txt','md') or raise ValueError."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed: "
            + ", ".join(sorted(ALLOWED_EXTENSIONS))
        )
    return ext.lstrip('.')


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes. Raises ValueError with human-readable message on failure."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            texts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            full_text = "\n".join(texts)
    except PDFPasswordIncorrect:
        raise ValueError("PDF is password-protected — remove password before uploading")
    except PDFSyntaxError:
        raise ValueError("PDF appears to be corrupt or is not a valid PDF file")
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")

    if not full_text.strip():
        raise ValueError(
            "PDF contains no extractable text — it may be a scanned image PDF. "
            "OCR is not supported."
        )
    return full_text


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes including table cells. Raises ValueError on corrupt file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except zipfile.BadZipFile:
        raise ValueError("DOCX file is corrupt or not a valid Word document")
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # MUST iterate tables separately — doc.paragraphs does not include table cells (OOXML structure)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ValueError("DOCX file contains no extractable text")
    return full_text


def parse_text(file_bytes: bytes) -> str:
    """Decode TXT or MD file bytes as UTF-8 (fallback: latin-1). Raises ValueError if empty."""
    try:
        full_text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        full_text = file_bytes.decode('latin-1', errors='replace')
    if not full_text.strip():
        raise ValueError("File contains no text content")
    return full_text


# Alias for backward compatibility with tests that call parse_txt directly
parse_txt = parse_text


# Explicit timeout to leave headroom within 60s Apache CGI limit.
# Default is undocumented (~20s); 15s + embed + overhead ~ 55s worst case (T-02-12).
_TRAFILATURA_TIMEOUT = 15


def fetch_and_extract_url(url: str) -> str:
    """Fetch URL and extract main text content via trafilatura.

    Raises ValueError with human-readable message if:
    - Network fetch fails (network error, SSL error, DNS failure)
    - Page contains no extractable text (JS-rendered, empty body, paywalled)

    Satisfies INGEST-04 and INGEST-07 (clear error for empty/JS pages).
    """
    # Set explicit download timeout before calling fetch_url.
    # trafilatura.settings.DOWNLOAD_TIMEOUT controls the urllib/requests timeout.
    trafilatura.settings.DOWNLOAD_TIMEOUT = _TRAFILATURA_TIMEOUT

    downloaded = trafilatura.fetch_url(url)
    if downloaded is None:
        raise ValueError(
            f"Failed to fetch URL — network error, SSL error, or timeout: {url}"
        )

    text = trafilatura.extract(downloaded)
    if text is None or not text.strip():
        raise ValueError(
            f"No extractable text found at URL — the page may be JavaScript-rendered "
            f"or contain no main content: {url}"
        )
    return text


def parse_file(file_bytes: bytes, filetype: str) -> str:
    """Dispatch to the correct parser based on filetype string ('pdf','docx','txt','md')."""
    if filetype == 'pdf':
        return parse_pdf(file_bytes)
    elif filetype == 'docx':
        return parse_docx(file_bytes)
    elif filetype in ('txt', 'md'):
        return parse_text(file_bytes)
    else:
        raise ValueError(f"No parser for filetype: {filetype}")
